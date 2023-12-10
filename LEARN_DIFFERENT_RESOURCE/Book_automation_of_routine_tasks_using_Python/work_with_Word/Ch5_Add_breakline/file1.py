import docx

doc = docx.Document()
doc.add_paragraph('Текст на первой странице')

doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('Текст на второй странице')
doc.save('twoPage.docx')