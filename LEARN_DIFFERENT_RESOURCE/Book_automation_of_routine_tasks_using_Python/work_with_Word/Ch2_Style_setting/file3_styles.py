import docx
doc = docx.Document('demo.docx')

print(doc.paragraphs[0].text)

print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'

print(doc.paragraphs[1].text)
print()

print(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text,
      doc.paragraphs[1].runs[3].text)

#doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].italic = True
doc.paragraphs[1].runs[3].bold = True
doc.paragraphs[1].runs[3].strike = True
doc.save('restyled.docx')
