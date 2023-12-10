import docx
doc = docx.Document()
doc.add_paragraph('Здравствуй, мир!')
paraObj1 = doc.add_paragraph('Это второй абзац.')
paraObj2 = doc.add_paragraph('Это еще один абзац.')
paraObj1.add_run('Этот текст добавляется во второй абзац')
doc.save('helloworld.docx')
