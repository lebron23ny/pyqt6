import docx
doc = docx.Document('helloworld.docx')
paragraphs_list = [paragraph.text for paragraph in doc.paragraphs]
print(paragraphs_list)


run_second_par = [run.text for run in doc.paragraphs[1].runs]
print(run_second_par)