import docx


doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
print(doc.paragraphs[0].style.name)
print(doc.paragraphs[1].style.name)
print(doc.paragraphs[2].style.name)
print(doc.paragraphs[3].style.name)
print(doc.paragraphs[4].style.name)
doc.save('headings.docx')