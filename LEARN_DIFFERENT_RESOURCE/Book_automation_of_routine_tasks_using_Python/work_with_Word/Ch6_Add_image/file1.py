import docx
doc = docx.Document('template.docx')
doc.add_paragraph('Здравствуй, мир!')

doc.add_picture('picture_3.png')
doc.add_paragraph('fgdd')

doc.paragraphs[2].style = 'Formula'
print(len(doc.paragraphs))

doc.save('document.docx')