import docx

doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
print()

print(doc.paragraphs[0].text)
print()
print(doc.paragraphs[1].text)
print()
print(len(doc.paragraphs[1].runs))
print()
print(doc.paragraphs[1].runs[0].text)
print()
print(doc.paragraphs[1].runs[1].text)
print()
print(doc.paragraphs[1].runs[2].text)
print()
print(doc.paragraphs[1].runs[3].text)
print()
t = doc.paragraphs[1].runs[4].text
print(doc.paragraphs[1].runs[4].text)


